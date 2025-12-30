"""
Text extraction pipeline for various document formats.
Supports PDF, DOCX, TXT, and images with OCR.
"""
import io
import os
import logging
from typing import List, Dict, Tuple, Optional
from pathlib import Path
import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)

class TextExtractor:
    """
    Text extraction pipeline that handles multiple document formats.
    """
    
    def __init__(self):
        # Set tesseract path if needed (adjust for your system)
        # pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'
        pass
    
    def extract_text(self, file_content: bytes, filename: str, mime_type: str) -> List[Dict[str, any]]:
        """
        Extract text from a file and return list of pages with text content.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: MIME type of the file
            
        Returns:
            List of dicts with page_num and text for each page
        """
        try:
            if mime_type == "application/pdf":
                return self._extract_pdf_text(file_content)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_docx_text(file_content)
            elif mime_type == "text/plain":
                return self._extract_txt_text(file_content)
            elif mime_type in ["image/png", "image/jpeg", "image/jpg", "image/gif", "image/bmp", "image/tiff", "image/webp"]:
                return self._extract_image_text(file_content)
            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                return []
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return []
    
    def _extract_pdf_text(self, file_content: bytes) -> List[Dict[str, any]]:
        """Extract text from PDF using both PyMuPDF and pdfplumber for best results."""
        pages = []
        
        try:
            # First try with pdfplumber for better text extraction
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        pages.append({
                            "page_num": page_num,
                            "text": text.strip()
                        })
                    else:
                        # If pdfplumber fails, try PyMuPDF
                        pages.append({
                            "page_num": page_num,
                            "text": ""
                        })
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyMuPDF")
            
            # Fallback to PyMuPDF
            try:
                pdf_doc = fitz.open(stream=file_content, filetype="pdf")
                pages = []
                for page_num in range(pdf_doc.page_count):
                    page = pdf_doc[page_num]
                    text = page.get_text()
                    pages.append({
                        "page_num": page_num + 1,
                        "text": text.strip()
                    })
                pdf_doc.close()
            except Exception as e2:
                logger.error(f"PyMuPDF also failed: {e2}")
                pages = []
        
        return pages
    
    def _extract_docx_text(self, file_content: bytes) -> List[Dict[str, any]]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            # Combine all text
            full_text = "\n".join(text_content).strip()
            
            return [{
                "page_num": 1,
                "text": full_text
            }] if full_text else []
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return []
    
    def _extract_txt_text(self, file_content: bytes) -> List[Dict[str, any]]:
        """Extract text from plain text file."""
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1', errors='ignore')
            
            return [{
                "page_num": 1,
                "text": text.strip()
            }] if text.strip() else []
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            return []
    
    def _extract_image_text(self, file_content: bytes) -> List[Dict[str, any]]:
        """Extract text from image using OCR."""
        try:
            # Open image with PIL
            image = Image.open(io.BytesIO(file_content))
            
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using tesseract
            text = pytesseract.image_to_string(image)
            
            return [{
                "page_num": 1,
                "text": text.strip()
            }] if text.strip() else []
            
        except Exception as e:
            logger.error(f"Error extracting image text: {e}")
            return []
    
    def extract_and_chunk_text(self, file_content: bytes, filename: str, mime_type: str, 
                              chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, any]]:
        """
        Extract text and split into chunks for embedding/retrieval.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: MIME type
            chunk_size: Target size for text chunks
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks with metadata
        """
        pages = self.extract_text(file_content, filename, mime_type)
        chunks = []
        
        for page in pages:
            text = page["text"]
            page_num = page["page_num"]
            
            # Simple text chunking
            if len(text) <= chunk_size:
                chunks.append({
                    "page_num": page_num,
                    "chunk_index": 0,
                    "text": text,
                    "start_char": 0,
                    "end_char": len(text)
                })
            else:
                # Split into overlapping chunks
                start = 0
                chunk_index = 0
                
                while start < len(text):
                    end = min(start + chunk_size, len(text))
                    
                    # Try to break at sentence boundaries
                    if end < len(text):
                        # Look for sentence end markers
                        for marker in ['. ', '.\n', '? ', '! ']:
                            last_marker = text.rfind(marker, start, end)
                            if last_marker != -1:
                                end = last_marker + len(marker)
                                break
                    
                    chunk_text = text[start:end].strip()
                    if chunk_text:
                        chunks.append({
                            "page_num": page_num,
                            "chunk_index": chunk_index,
                            "text": chunk_text,
                            "start_char": start,
                            "end_char": end
                        })
                    
                    # Move start position with overlap
                    start = max(start + chunk_size - overlap, end)
                    chunk_index += 1
                    
                    if start >= len(text):
                        break
        
        return chunks

# Global extractor instance
text_extractor = TextExtractor()